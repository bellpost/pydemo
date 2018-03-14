#!/usr/bin/python
# coding=utf-8
'''
@ctime: 2017年3月22日
@author: Bellpost
'''
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText

from collections import defaultdict
from datetime import datetime
import time,sys
from bs4 import BeautifulSoup 
import requests
from docx import Document

class getBusNews_wh():
    def __init__(self,conf_file,doc_name):
        self.document = Document()
        self.conf_file = conf_file
        self.doc_name=doc_name
        
    #获取本地存储时间
    def getLocalDate(self):
        try:
            with open(conf_file,'rb') as rf:
                tranDate = rf.readline()
                year_s, mon_s, day_s = tranDate.split('-')
            return datetime(int(year_s), int(mon_s), int(day_s))
        except Exception, e:  
            print "Read localtime failed:"+str(e)
            sys.exit(1)
    
    #写入当前更新end时间
    def writeDate(self,curDate):
        try:
            with open(conf_file,'wb+') as wf:
                wf.write(curDate.strftime('%Y-%m-%d'))
                wf.close()
        except Exception, e:  
            print "write new time failed:"+str(e)
            with open(conf_file,'wb+') as wf:
                wf.write(time.strftime('%Y-%m-%d'))
                wf.close()
            print "write current time successfully!"
            sys.exit(1)
    
    
    #获取网页内容
    def get_pageData(self,url):
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_11_2) '
                          'AppleWebKit/537.36 (KHTML, like Gecko) Chrome/47.0.2526.80 Safari/537.36'
        }
        data = requests.get(url, headers=headers).content    
        return data
    
    #格式化时间
    def parseTime(self,timeStr):
        year_s, mon_s, day_s=timeStr.split('-')
        return datetime(int(year_s), int(mon_s), int(day_s))
       
    #获取官网更新的最新的时间
    def parse_NewsDate(self,html):
        soup = BeautifulSoup(html, 'html.parser')
        #print soup  
        tag_date = soup.find('font', attrs={'color': '#BEBEBE'}).getText()
        news_Date = tag_date.replace("[","").replace("]","")
        return self.parseTime(news_Date)
    
    
    #根据获取数据
    def getContent(self,html):
        soup = BeautifulSoup(html, 'html.parser')  
        try:
            title = soup.find('div', attrs={'class': 'btitilef'}).getText()
        except Exception, e:  
            title = "title is null"
            print "Title Error"+str(e)
        try:
            daytime = soup.find('div', attrs={'class': 'btdaytime'}).getText()
        except Exception, e:  
            daytime = "daytime is null"
            print "Title Error"+str(e)
        try:
            conttxt = soup.find('div', attrs={'id': 'Conttxt'})
        except Exception, e:  
            conttxt = "conttxt is null"
            print "Title Error"+str(e)
        paras = []
        try:
            for paragraph in conttxt.find_all('p'):
                p_content = paragraph.getText()
                paras.append(p_content)
        except Exception, e:  
            paras = "paras is null"
            print "paras Error"+str(e)
        return title,daytime,paras
    
    #获取时间和连接
    def parse_NewsDateandUrl(self,html):
        soup = BeautifulSoup(html, 'html.parser')
        contonedate = soup.findAll("font")
        page=[]
        n=0
        dict = defaultdict(lambda: 'none')
        for link in soup.findAll("a",attrs={'style':"color:#19a3f1"}):
            page.append(link)
        for i in contonedate:
            i=i.getText().replace("[","").replace("]","")
            dict[n] = (i,page[n].attrs["href"])
            n=n+1
        return dict   
    
    #写入docx
    def page_docx(self,title,daytime,cont):
        self.document.add_heading(title,2)
        self.document.add_heading(daytime,3)
        for p in cont:
            self.document.add_paragraph('%s\n' % p.strip().replace("\n",""))
            
    def save_docx(self):       
        self.document.save(self.doc_name)
        print 'saves docx file：',self.doc_name

                
        

conf_file='date.conf'
index_url = "http://www.wuhanbus.com/html/class/xlxx/index.html"
initcont_url = "http://www.wuhanbus.com"
date_title = time.strftime('%Y%m%d%H')
doc_name = 'wuhanbusnews_daily'+date_title+'.docx'
HOST = "smtp.139.com"
SUBJECT = u"<Auto>武汉公交监测程序和官网新闻"+date_title
TO = "--"
FROM = "-@139.com"

def mail_cont():
    msg = MIMEMultipart('related')
    msgtext = MIMEText("<font color=red>武汉公交官网公交更新详细内容见附件。</font>","html","utf-8")
    msg.attach(msgtext)
    attach = MIMEText(open(doc_name, "rb").read(), "base64", "utf-8")
    attach["Content-Type"] = "application/octet-stream"
    attach["Content-Disposition"] = "attachment; filename=\"武汉公交日报.docx\""
    msg.attach(attach)
    msg['Subject'] = SUBJECT
    msg['From']=FROM
    msg['To']=TO
    return msg

#获取相等的日期
def indexurl():
    index_urls=[]
    for index in range(1,8):
        if index == 1:
            index_url = "http://www.wuhanbus.com/html/class/xlxx/index.html"
            index_urls.append(index_url)
        else:
            index_url = "http://www.wuhanbus.com/html/class/xlxx/index_"+str(index)+".html"
            index_urls.append(index_url)
    return index_urls

#获取历史数据
def main1():
    index_url1 = "http://www.wuhanbus.com/html/class/xlxx/index.html"
    wuhan=getBusNews_wh(conf_file,doc_name)
    data = wuhan.get_pageData(index_url1)
    olddate = wuhan.getLocalDate()
    newdate = wuhan.parse_NewsDate(data)
    n=0
    dict = defaultdict(lambda: 'none')
    for index_url in indexurl():
        data = wuhan.get_pageData(index_url)
        soup = BeautifulSoup(data, 'html.parser')
        contonedate = soup.findAll("font")
        page=[]
        m=contonedate.__len__()
        for link in soup.findAll("a",attrs={'style':"color:#19a3f1"}):
            page.append(link)
        for i in contonedate:
            if m == contonedate.__len__():m=0
            i=i.getText().replace("[","").replace("]","")
            dict[n] = (i,page[m].attrs["href"])
            n=n+1
            m=m+1
    #print dict

    if newdate > olddate:
        for i in range(dict.__len__()):
            datelistone,urlone = dict[i]
            conturl = initcont_url + urlone
            #print conturl
            if wuhan.parseTime(datelistone) > olddate:
                title,daytime,cont=wuhan.getContent(wuhan.get_pageData(conturl))
                wuhan.page_docx(title,daytime,cont)
        wuhan.save_docx()
        wuhan.writeDate(newdate)
        
        msg = mail_cont()
        while True:
            try:
                print "-------------------------------"
                server = smtplib.SMTP()
                server.connect(HOST,"25")
                #server.starttls()
                server.login("13997578413@139.com","zbw19941217")
                server.sendmail(FROM, TO, msg.as_string())
                print "Email send Successful!"
                server.quit()
                sys.exit(0)
            except Exception, e:  
                print "Email send failed:"+str(e)
                continue

    
def main():
    wuhan=getBusNews_wh(conf_file,doc_name)
    data = wuhan.get_pageData(index_url)
    olddate = wuhan.getLocalDate()
    newdate = wuhan.parse_NewsDate(data)
    dict = wuhan.parse_NewsDateandUrl(data)
    if newdate > olddate:
        for i in range(dict.__len__()):
            datelistone,urlone = dict[i]
            conturl = initcont_url + urlone
            #print conturl
            if wuhan.parseTime(datelistone) > olddate:
                title,daytime,cont=wuhan.getContent(wuhan.get_pageData(conturl))
                wuhan.page_docx(title,daytime,cont)
        wuhan.save_docx()
        wuhan.writeDate(newdate)
        
        msg = mail_cont()
        while True:
            try:
                server = smtplib.SMTP()
                server.connect(HOST,"25")
                server.login("-@139.com","--")
                server.sendmail(FROM, TO, msg.as_string())
                print "Email send Successful!"
                server.quit()
                sys.exit(0)
            except Exception, e:  
                print "Email send failed:"+str(e)
                continue
    
if __name__=="__main__":
    main1()
    